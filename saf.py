from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import subprocess

def add_hyperlink(paragraph, text, url):
    """Adds a hyperlink to a paragraph with blue color and underline."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = parse_xml(r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="%s"/>' % r_id)
    
    new_run = paragraph.add_run()
    new_run._r.append(hyperlink)
    
    # Set the text and format it to look like a hyperlink
    new_run.text = text
    new_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
    new_run.font.underline = True
    
    return hyperlink

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Environmental & Safety Engineering 4-Year Ghana-Specific Career Roadmap', 0)

# Introduction
intro = doc.add_paragraph()
intro.add_run('Tailored for Ghanaian Students - Focusing on Local Regulations, Industries, and Opportunities\n\n').bold = True
intro.add_run(
    "This roadmap is specifically designed for Environmental & Safety Engineering students in Ghana, "
    "incorporating local certifications, Ghanaian regulations, key industries, and essential skill development."
)

# Year 1: Foundations & Ghana Context
doc.add_heading('Year 1: Ghanaian Context & Foundations', level=1)
year1 = doc.add_paragraph()
year1.add_run('Core Focus:\n').bold = True
year1.add_run(
    "• Establish foundation in Ghana's environmental and safety landscape\n"
    "• Understand key Ghanaian regulations and agencies\n"
    "• Develop basic technical skills with local industry focus\n\n"
)

year1.add_run('Ghana-Specific Certifications:\n').bold = True
year1.add_run("• ")
add_hyperlink(year1, "EPA Ghana Basic Environmental Awareness Certificate", "https://www.epa.gov.gh")
year1.add_run("\n• ")
add_hyperlink(year1, "Minerals Commission Basic Mine Safety Induction", "https://www.mincom.gov.gh")
year1.add_run("\n• ")
add_hyperlink(year1, "Ghana Red Cross Society First Aid & Emergency Response", "https://www.redcross.org.gh")
year1.add_run("\n• ")
add_hyperlink(year1, "Fire Service Basic Fire Safety Certificate", "https://www.ghananationalfireervice.gov.gh")
year1.add_run("\n\n")

year1.add_run('Key Ghanaian Regulations to Master:\n').bold = True
year1.add_run(
    "• Environmental Protection Agency Act, 1994 (Act 490)\n"
    "• Minerals and Mining Act, 2006 (Act 703)\n"
    "• Factories, Offices and Shops Act, 1970 (Act 328)\n\n"
)

year1.add_run('Essential Skills to Develop:\n').bold = True
year1.add_run(
    "• Technical Writing: Learn to write clear environmental and safety reports\n"
    "• Basic Data Analysis: Excel for environmental data tracking and basic statistics\n"
    "• Communication: Present technical information clearly to non-technical audiences\n"
    "• Problem-Solving: Basic root cause analysis for safety incidents\n"
    "• Time Management: Balance academic workload with certification studies\n"
    "• Digital Literacy: Microsoft Office suite proficiency\n"
    "• Laboratory Skills: Basic chemical handling and safety procedures\n"
    "• Regulatory Navigation: Understanding how to find and interpret Ghanaian laws\n"
)

# Year 2: Technical Skills & Ghanaian Standards
doc.add_heading('Year 2: Ghana Standards & Technical Applications', level=1)
year2 = doc.add_paragraph()
year2.add_run('Core Focus:\n').bold = True
year2.add_run(
    "• Apply engineering principles to Ghana's key industries\n"
    "• Develop proficiency with local environmental standards\n"
    "• Begin industry exposure through local site visits\n\n"
)

year2.add_run('Intermediate Ghana Certifications:\n').bold = True
year2.add_run("• ")
add_hyperlink(year2, "EPA Ghana Environmental Impact Assessment (EIA) Procedures", "https://www.epa.gov.gh/epa/eia-division/")
year2.add_run("\n• ")
add_hyperlink(year2, "Ghana Standards Authority (GSA) Quality Systems Training", "https://www.gsa.gov.gh")
year2.add_run("\n• ")
add_hyperlink(year2, "Minerals Commission Advanced Safety Certification", "https://www.mincom.gov.gh")
year2.add_run("\n• ")
add_hyperlink(year2, "Ghana Water Company Wastewater Management Basics", "https://www.gwcl.com.gh")
year2.add_run("\n\n")

year2.add_run('Ghana Industry Software Skills:\n').bold = True
year2.add_run(
    "• AutoCAD for mining and construction layouts\n"
    "• GIS applications for environmental mapping in Ghana\n"
    "• Excel for Ghana EPA compliance reporting\n"
    "• Basic programming for environmental data analysis\n\n"
)

year2.add_run('Local Industry Exposure:\n').bold = True
year2.add_run(
    "• Site visits to local mines (Tarkwa, Obuasi)\n"
    "• Manufacturing plant tours (Accra, Tema Industrial Area)\n"
    "• Water treatment plant visits\n"
    "• Oil & gas facility orientation (Takoradi)\n\n"
)

year2.add_run('Advanced Skills to Develop:\n').bold = True
year2.add_run(
    "• Risk Assessment: Conduct basic job safety analysis and environmental risk assessments\n"
    "• Technical Drawing: Interpret and create basic engineering drawings\n"
    "• Data Visualization: Create charts and graphs for environmental monitoring data\n"
    "• Project Management: Basic project planning and timeline management\n"
    "• Stakeholder Engagement: Interacting with community members and regulatory officials\n"
    "• Research Skills: Literature review and technical research methods\n"
    "• Quality Assurance: Understanding quality control processes in Ghanaian industries\n"
    "• Environmental Monitoring: Basic air, water, and soil sampling techniques\n"
    "• Safety Auditing: Conduct basic workplace safety inspections\n"
)

# Year 3: Specialization & Ghana Industry Integration
doc.add_heading('Year 3: Ghana Industry Specialization', level=1)
year3 = doc.add_paragraph()
year3.add_run('Core Focus:\n').bold = True
year3.add_run(
    "• Specialize in Ghana's priority sectors\n"
    "• Gain practical experience through industrial attachment\n"
    "• Develop risk assessment skills for local contexts\n\n"
)

year3.add_run('Specialization Tracks (Choose based on Ghana market):\n').bold = True
year3.add_run(
    "Mining & Minerals Track:\n"
    "• Mine Safety & Emergency Response\n"
    "• Tailings Dam Management\n"
    "• Cyanide Management Code (for gold mining)\n\n"
)

year3.add_run("Oil & Gas Track:\n").bold = True
year3.add_run(
    "• Offshore Safety Procedures\n"
    "• Petroleum Industry HSE Standards\n"
    "• Spill Prevention & Response\n\n"
)

year3.add_run("Manufacturing & Construction Track:\n").bold = True
year3.add_run(
    "• Factory Act Compliance\n"
    "• Construction Site Safety (Ghana context)\n"
    "• Industrial Waste Management\n\n"
)

year3.add_run('Ghana Industrial Attachment:\n').bold = True
year3.add_run(
    "• Summer internship with Ghanaian companies:\n"
    "  - "
)
add_hyperlink(year3, "Gold Fields Ghana", "https://www.goldfields.com")
year3.add_run(", ")
add_hyperlink(year3, "Newmont Ghana", "https://www.newmont.com")
year3.add_run(", ")
add_hyperlink(year3, "Anglogold Ashanti", "https://www.anglogoldashanti.com")
year3.add_run("\n  - ")
add_hyperlink(year3, "Tullow Ghana", "https://www.tullowoil.com")
year3.add_run(", ")
add_hyperlink(year3, "GNPC", "https://www.gnpcghana.com")
year3.add_run(", ")
add_hyperlink(year3, "GOIL", "https://www.goil.com.gh")
year3.add_run("\n  - ")
add_hyperlink(year3, "Unilever Ghana", "https://www.unilever-ghana.com")
year3.add_run(", ")
add_hyperlink(year3, "Nestlé Ghana", "https://www.nestle-ghana.com")
year3.add_run(", ")
add_hyperlink(year3, "Guinness Ghana", "https://www.guinnessghana.com")
year3.add_run("\n  - Construction firms (Mansco, Consar, etc.)\n\n")

year3.add_run('Professional Skills to Master:\n').bold = True
year3.add_run(
    "• Advanced Risk Management: Quantitative risk assessment and bow-tie analysis\n"
    "• Incident Investigation: Root cause analysis using methodologies like 5-Whys\n"
    "• Environmental Management Systems: ISO 14001 implementation and auditing\n"
    "• Safety Leadership: Influencing safety culture and behavior-based safety\n"
    "• Technical Reporting: Writing comprehensive EIA reports and safety cases\n"
    "• Budget Management: Cost estimation for safety and environmental projects\n"
    "• Regulatory Compliance: Navigating complex multi-agency requirements\n"
    "• Emergency Response Planning: Developing and testing emergency procedures\n"
    "• Contract Management: Understanding contractor safety management\n"
    "• Cultural Competence: Working effectively in Ghana's diverse work environments\n"
)

# Year 4: Professional Integration & Career Launch
doc.add_heading('Year 4: Ghana Professional Integration', level=1)
year4 = doc.add_paragraph()
year4.add_run('Core Focus:\n').bold = True
year4.add_run(
    "• Finalize professional certifications\n"
    "• Conduct Ghana-focused research project\n"
    "• Transition to employment in Ghanaian industries\n\n"
)

year4.add_run('Advanced Ghana Certifications:\n').bold = True
year4.add_run("• ")
add_hyperlink(year4, "EPA Ghana Environmental Inspector Preparation", "https://www.epa.gov.gh")
year4.add_run("\n• ISO 14001:2015 (Environmental Management) - Local auditors\n")
year4.add_run("• ISO 45001:2018 (Occupational Health & Safety) - Local context\n")
year4.add_run("• ")
add_hyperlink(year4, "NEBOSH International Diploma", "https://www.nebosh.org.uk")
year4.add_run(" (if resources allow)\n\n")

year4.add_run('Final Year Project (Ghana Focus):\n').bold = True
year4.add_run(
    "• Environmental impact of galamsey (illegal mining)\n"
    "• Safety systems in Ghana's oil & gas industry\n"
    "• Waste management solutions for Ghanaian cities\n"
    "• Industrial pollution control in Ghana\n"
    "• Renewable energy safety standards for Ghana\n\n"
)

year4.add_run('Career Preparation - Ghana Market:\n').bold = True
year4.add_run("• Join ")
add_hyperlink(year4, "Ghana Institution of Engineers (GhIE)", "https://www.ghie.org.gh")
year4.add_run("\n• Register with Ghana Institution of Safety and Environment Professionals\n")
year4.add_run("• Attend Ghana Mining Industry career fairs\n")
year4.add_run("• Prepare for Ghanaian employer expectations\n")
year4.add_run("• Network at Ghana Oil & Gas conferences\n\n")

year4.add_run('Leadership & Strategic Skills:\n').bold = True
year4.add_run(
    "• Strategic Planning: Developing departmental safety and environmental strategies\n"
    "• Change Management: Implementing new safety systems and procedures\n"
    "• Financial Acumen: Budgeting and cost-benefit analysis for HSE projects\n"
    "• Negotiation Skills: Dealing with regulators, contractors, and stakeholders\n"
    "• Crisis Management: Leading during environmental or safety emergencies\n"
    "• Mentorship: Training and developing junior staff and technicians\n"
    "• Public Speaking: Presenting to senior management and regulatory bodies\n"
    "• Business Development: Contributing to bids and proposals with HSE components\n"
    "• Continuous Improvement: Implementing Kaizen and other improvement methodologies\n"
    "• Digital Transformation: Leveraging technology for HSE management systems\n"
)

# Key Ghanaian Organizations & Contacts
doc.add_heading('Essential Ghanaian Organizations & Resources', level=1)
orgs = doc.add_paragraph()
orgs.add_run('Regulatory Bodies:\n').bold = True
orgs.add_run("• ")
add_hyperlink(orgs, "Environmental Protection Agency (EPA) Ghana", "https://www.epa.gov.gh")
orgs.add_run("\n• ")
add_hyperlink(orgs, "Minerals Commission of Ghana", "https://www.mincom.gov.gh")
orgs.add_run("\n• ")
add_hyperlink(orgs, "Ghana Standards Authority", "https://www.gsa.gov.gh")
orgs.add_run("\n• Factories Inspectorate Department\n")
orgs.add_run("• ")
add_hyperlink(orgs, "National Fire Service", "https://www.ghananationalfireervice.gov.gh")
orgs.add_run("\n\n")

orgs.add_run('Professional Associations:\n').bold = True
orgs.add_run("• ")
add_hyperlink(orgs, "Ghana Institution of Engineers (GhIE)", "https://www.ghie.org.gh")
orgs.add_run("\n• Ghana Institution of Safety and Environment Professionals\n")
orgs.add_run("• ")
add_hyperlink(orgs, "Ghana Mining Society", "https://www.ghanaminingsociety.org")
orgs.add_run("\n• ")
add_hyperlink(orgs, "Association of Ghana Industries", "https://www.agighana.org")
orgs.add_run("\n\n")

orgs.add_run('Key Industries for Employment:\n').bold = True
orgs.add_run("• Mining: ")
add_hyperlink(orgs, "Newmont", "https://www.newmont.com")
orgs.add_run(", ")
add_hyperlink(orgs, "Gold Fields", "https://www.goldfields.com")
orgs.add_run(", ")
add_hyperlink(orgs, "Anglogold Ashanti", "https://www.anglogoldashanti.com")
orgs.add_run(", Golden Star\n")
orgs.add_run("• Oil & Gas: ")
add_hyperlink(orgs, "Tullow", "https://www.tullowoil.com")
orgs.add_run(", ")
add_hyperlink(orgs, "GNPC", "https://www.gnpcghana.com")
orgs.add_run(", ")
add_hyperlink(orgs, "GOIL", "https://www.goil.com.gh")
orgs.add_run(", Springfield, ENI\n")
orgs.add_run("• Manufacturing: ")
add_hyperlink(orgs, "Unilever", "https://www.unilever-ghana.com")
orgs.add_run(", ")
add_hyperlink(orgs, "Nestlé", "https://www.nestle-ghana.com")
orgs.add_run(", ")
add_hyperlink(orgs, "Guinness", "https://www.guinnessghana.com")
orgs.add_run(", FanMilk, Cocoa Processing\n")
orgs.add_run("• Construction: Mansco, Consar, Maripoma, Engineers & Planners\n")
orgs.add_run("• Utilities: ")
add_hyperlink(orgs, "Ghana Water Company", "https://www.gwcl.com.gh")
orgs.add_run(", ECG, VRA\n")

# Success Tips for Ghanaian Context
doc.add_heading('Success Strategies for Ghanaian Graduates', level=1)
tips = doc.add_paragraph()
tips.add_run('Academic Excellence:\n').bold = True
tips.add_run(
    "• Maintain strong GPA (minimum 3.0 for competitive positions)\n"
    "• Develop strong technical writing skills for Ghanaian reports\n"
    "• Master Ghana's environmental regulations and standards\n\n"
)

tips.add_run('Professional Development:\n').bold = True
tips.add_run(
    "• Start with basic Ghana EPA certifications in Year 1\n"
    "• Build relationships with Ghanaian professionals early\n"
    "• Attend Ghana-specific industry workshops and seminars\n"
    "• Develop understanding of Ghanaian business culture\n\n"
)

tips.add_run('Networking in Ghana:\n').bold = True
tips.add_run("• Join ")
add_hyperlink(tips, "GhIE student chapters", "https://www.ghie.org.gh")
tips.add_run("\n• Attend Ghana Mining Industry events\n")
tips.add_run("• Connect with alumni working in Ghanaian industries\n")
tips.add_run("• Participate in Ghana Environmental Protection forums\n")

# Skill Development Summary
doc.add_heading('4-Year Skill Development Progression', level=1)
skills_summary = doc.add_paragraph()
skills_summary.add_run('Year 1 - Foundation Skills:\n').bold = True
skills_summary.add_run("Technical Writing, Basic Data Analysis, Communication, Problem-Solving, Time Management\n\n")

skills_summary.add_run('Year 2 - Technical Skills:\n').bold = True
skills_summary.add_run("Risk Assessment, Technical Drawing, Data Visualization, Project Management, Stakeholder Engagement\n\n")

skills_summary.add_run('Year 3 - Professional Skills:\n').bold = True
skills_summary.add_run("Advanced Risk Management, Incident Investigation, EMS Implementation, Safety Leadership, Regulatory Compliance\n\n")

skills_summary.add_run('Year 4 - Leadership Skills:\n').bold = True
skills_summary.add_run("Strategic Planning, Change Management, Financial Acumen, Crisis Management, Business Development\n")

# Save the document to current directory
file_path = "Environmental_Safety_Engineering_Ghana_Roadmap.docx"
doc.save(file_path)

print("✅ Comprehensive Ghana-specific roadmap with working hyperlinks created successfully!")
print(f"📁 File saved at: {file_path}")
