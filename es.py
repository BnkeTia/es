from docx import Document
from docx.shared import Inches
import os
import subprocess
import getpass

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Environmental & Safety Engineering 4-Year Ghana-Specific Career Roadmap', 0)

# Introduction
intro = doc.add_paragraph()
intro.add_run('Tailored for Ghanaian Students - Focusing on Local Regulations, Industries, and Opportunities\n\n').bold = True
intro.add_run(
    "This roadmap is specifically designed for Environmental & Safety Engineering students in Ghana, "
    "incorporating local certifications, Ghanaian regulations, and key industries that drive the national economy."
)

# Year 1: Foundations & Ghana Context
doc.add_heading('Year 1: Ghanaian Context & Foundations', level=1)
year1 = doc.add_paragraph()
year1.add_run('Core Focus:\n').bold = True
year1.add_run(
    "• Establish foundation in Ghana's environmental and safety landscape\n"
    "• Understand key Ghanaian regulations and agencies\n"
    "• Develop basic technical skills with local industry focus\n\n"
)

year1.add_run('Ghana-Specific Certifications:\n').bold = True
year1.add_run(
    "• EPA Ghana Basic Environmental Awareness Certificate\n"
    "• Minerals Commission Basic Mine Safety Induction\n"
    "• Ghana Red Cross Society First Aid & Emergency Response\n"
    "• Fire Service Basic Fire Safety Certificate\n\n"
)

year1.add_run('Key Ghanaian Regulations to Master:\n').bold = True
year1.add_run(
    "• Environmental Protection Agency Act, 1994 (Act 490)\n"
    "• Minerals and Mining Act, 2006 (Act 703)\n"
    "• Factories, Offices and Shops Act, 1970 (Act 328)\n"
)

# Year 2: Technical Skills & Ghanaian Standards
doc.add_heading('Year 2: Ghana Standards & Technical Applications', level=1)
year2 = doc.add_paragraph()
year2.add_run('Core Focus:\n').bold = True
year2.add_run(
    "• Apply engineering principles to Ghana's key industries\n"
    "• Develop proficiency with local environmental standards\n"
    "• Begin industry exposure through local site visits\n\n"
)

year2.add_run('Intermediate Ghana Certifications:\n').bold = True
year2.add_run(
    "• EPA Ghana Environmental Impact Assessment (EIA) Procedures\n"
    "• Ghana Standards Authority (GSA) Quality Systems Training\n"
    "• Minerals Commission Advanced Safety Certification\n"
    "• Ghana Water Company Wastewater Management Basics\n\n"
)

year2.add_run('Ghana Industry Software Skills:\n').bold = True
year2.add_run(
    "• AutoCAD for mining and construction layouts\n"
    "• GIS applications for environmental mapping in Ghana\n"
    "• Excel for Ghana EPA compliance reporting\n"
    "• Basic programming for environmental data analysis\n\n"
)

year2.add_run('Local Industry Exposure:\n').bold = True
year2.add_run(
    "• Site visits to local mines (Tarkwa, Obuasi)\n"
    "• Manufacturing plant tours (Accra, Tema Industrial Area)\n"
    "• Water treatment plant visits\n"
    "• Oil & gas facility orientation (Takoradi)\n"
)

# Year 3: Specialization & Ghana Industry Integration
doc.add_heading('Year 3: Ghana Industry Specialization', level=1)
year3 = doc.add_paragraph()
year3.add_run('Core Focus:\n').bold = True
year3.add_run(
    "• Specialize in Ghana's priority sectors\n"
    "• Gain practical experience through industrial attachment\n"
    "• Develop risk assessment skills for local contexts\n\n"
)

year3.add_run('Specialization Tracks (Choose based on Ghana market):\n').bold = True
year3.add_run(
    "Mining & Minerals Track:\n"
    "• Mine Safety & Emergency Response\n"
    "• Tailings Dam Management\n"
    "• Cyanide Management Code (for gold mining)\n\n"
)

year3.add_run("Oil & Gas Track:\n").bold = True
year3.add_run(
    "• Offshore Safety Procedures\n"
    "• Petroleum Industry HSE Standards\n"
    "• Spill Prevention & Response\n\n"
)

year3.add_run("Manufacturing & Construction Track:\n").bold = True
year3.add_run(
    "• Factory Act Compliance\n"
    "• Construction Site Safety (Ghana context)\n"
    "• Industrial Waste Management\n\n"
)

year3.add_run('Ghana Industrial Attachment:\n').bold = True
year3.add_run(
    "• Summer internship with Ghanaian companies:\n"
    "  - Gold Fields Ghana, Newmont Ghana, Anglogold Ashanti\n"
    "  - Tullow Ghana, GNPC, GOIL\n"
    "  - Unilever Ghana, Nestlé Ghana, Guinness Ghana\n"
    "  - Construction firms (Mansco, Consar, etc.)\n"
)

# Year 4: Professional Integration & Career Launch
doc.add_heading('Year 4: Ghana Professional Integration', level=1)
year4 = doc.add_paragraph()
year4.add_run('Core Focus:\n').bold = True
year4.add_run(
    "• Finalize professional certifications\n"
    "• Conduct Ghana-focused research project\n"
    "• Transition to employment in Ghanaian industries\n\n"
)

year4.add_run('Advanced Ghana Certifications:\n').bold = True
year4.add_run(
    "• EPA Ghana Environmental Inspector Preparation\n"
    "• ISO 14001:2015 (Environmental Management) - Local auditors\n"
    "• ISO 45001:2018 (Occupational Health & Safety) - Local context\n"
    "• NEBOSH International Diploma (if resources allow)\n\n"
)

year4.add_run('Final Year Project (Ghana Focus):\n').bold = True
year4.add_run(
    "• Environmental impact of galamsey (illegal mining)\n"
    "• Safety systems in Ghana's oil & gas industry\n"
    "• Waste management solutions for Ghanaian cities\n"
    "• Industrial pollution control in Ghana\n"
    "• Renewable energy safety standards for Ghana\n\n"
)

year4.add_run('Career Preparation - Ghana Market:\n').bold = True
year4.add_run(
    "• Join Ghana Institution of Engineers (GhIE)\n"
    "• Register with Ghana Institution of Safety and Environment Professionals\n"
    "• Attend Ghana Mining Industry career fairs\n"
    "• Prepare for Ghanaian employer expectations\n"
    "• Network at Ghana Oil & Gas conferences\n"
)

# Key Ghanaian Organizations & Contacts
doc.add_heading('Essential Ghanaian Organizations & Resources', level=1)
orgs = doc.add_paragraph()
orgs.add_run('Regulatory Bodies:\n').bold = True
orgs.add_run(
    "• Environmental Protection Agency (EPA) Ghana\n"
    "• Minerals Commission of Ghana\n"
    "• Ghana Standards Authority\n"
    "• Factories Inspectorate Department\n"
    "• National Fire Service\n\n"
)

orgs.add_run('Professional Associations:\n').bold = True
orgs.add_run(
    "• Ghana Institution of Engineers (GhIE)\n"
    "• Ghana Institution of Safety and Environment Professionals\n"
    "• Ghana Mining Society\n"
    "• Association of Ghana Industries\n\n"
)

orgs.add_run('Key Industries for Employment:\n').bold = True
orgs.add_run(
    "• Mining: Newmont, Gold Fields, Anglogold Ashanti, Golden Star\n"
    "• Oil & Gas: Tullow, GNPC, GOIL, Springfield, ENI\n"
    "• Manufacturing: Unilever, Nestlé, Guinness, FanMilk, Cocoa Processing\n"
    "• Construction: Mansco, Consar, Maripoma, Engineers & Planners\n"
    "• Utilities: Ghana Water Company, ECG, VRA\n"
)

# Success Tips for Ghanaian Context
doc.add_heading('Success Strategies for Ghanaian Graduates', level=1)
tips = doc.add_paragraph()
tips.add_run('Academic Excellence:\n').bold = True
tips.add_run(
    "• Maintain strong GPA (minimum 3.0 for competitive positions)\n"
    "• Develop strong technical writing skills for Ghanaian reports\n"
    "• Master Ghana's environmental regulations and standards\n\n"
)

tips.add_run('Professional Development:\n').bold = True
tips.add_run(
    "• Start with basic Ghana EPA certifications in Year 1\n"
    "• Build relationships with Ghanaian professionals early\n"
    "• Attend Ghana-specific industry workshops and seminars\n"
    "• Develop understanding of Ghanaian business culture\n\n"
)

tips.add_run('Networking in Ghana:\n').bold = True
tips.add_run(
    "• Join GhIE student chapters\n"
    "• Attend Ghana Mining Industry events\n"
    "• Connect with alumni working in Ghanaian industries\n"
    "• Participate in Ghana Environmental Protection forums\n"
)

# Save the document to current directory
file_path = "Environmental_Safety_Engineering_Ghana_Roadmap.docx"
doc.save(file_path)

print("Ghana-specific roadmap created successfully!")
print(f"File saved at: {file_path}")

# Git operations
try:
    # Initialize git repository
    print("\nInitializing Git repository...")
    subprocess.run(["git", "init", "es"], check=True)
    os.chdir("es")
    
    # Configure git (basic configuration)
    subprocess.run(["git", "config", "user.name", "ES Roadmap Creator"], check=True)
    subprocess.run(["git", "config", "user.email", "es-roadmap@example.com"], check=True)
    
    # Move the created file to the repository
    import shutil
    shutil.move(f"../{file_path}", ".")
    
    # Add file to git
    print("Adding file to Git...")
    subprocess.run(["git", "add", file_path], check=True)
    
    # Commit the file
    print("Committing file...")
    subprocess.run(["git", "commit", "-m", "Initial commit: Ghana Environmental & Safety Engineering Roadmap"], check=True)
    
    print(f"\n✅ Success! Git repository 'es' initialized and file committed.")
    print(f"📁 Repository location: {os.getcwd()}")
    print(f"📄 File committed: {file_path}")
    
    # Instructions for pushing to GitHub
    print("\n" + "="*60)
    print("To push to GitHub, you need to:")
    print("1. Create a new repository on GitHub named 'es'")
    print("2. Run these commands:")
    print(f"   cd {os.getcwd()}")
    print('   git remote add origin https://github.com/YOUR_USERNAME/es.git')
    print('   git branch -M main')
    print('   git push -u origin main')
    print("="*60)
    
except subprocess.CalledProcessError as e:
    print(f"❌ Git error: {e}")
    print("Make sure Git is installed and configured on your system.")
except Exception as e:
    print(f"❌ An error occurred: {e}")
